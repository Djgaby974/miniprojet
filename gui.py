import tkinter as tk
import customtkinter as ctk
import ollama
import threading
import re
import os
import json
import pyperclip
import markdown2
from tkinter import filedialog, messagebox
from typing import List, Dict, Any
import random
import PyPDF2
import docx
import textract
import requests
from PIL import Image
from io import BytesIO
import queue
import time
import webbrowser
from bs4 import BeautifulSoup
import urllib.parse
import pdfplumber

# Palette de couleurs (conservée de l'ancienne implémentation)
COULEURS = {
    "fond_principal": "#1E1E2E",
    "accent_primaire": "#7E57C2",
    "accent_secondaire": "#4DB6AC",
    "texte_principal": "#E0E0E0",
    "texte_secondaire": "#B0BEC5",
    "fond_chat": "#2C3E50",
    "fond_code": "#263238",
    "fond_secondaire": "#2C3E50"
}

class ConversationOnglet(ctk.CTkFrame):
    def __init__(self, master, gestionnaire_config, gestionnaire_ollama, **kwargs):
        super().__init__(
            master, 
            fg_color=COULEURS["fond_principal"],  
            corner_radius=10,
            **kwargs
        )
        
        # Stocker le gestionnaire Ollama
        self.gestionnaire_ollama = gestionnaire_ollama
        
        self.gestionnaire_config = gestionnaire_config
        self.contexte_actuel = "Développement"
        self.fichier_charge = None
        self.dernier_chemin_fichier = None
        
        # Cadre de chat avec style IDE
        self.cadre_chat = ctk.CTkScrollableFrame(
            self, 
            fg_color=COULEURS["fond_chat"],  
            corner_radius=10
        )
        self.cadre_chat.pack(expand=True, fill='both', padx=10, pady=10)
        
        # Cadre de saisie
        self.cadre_saisie = ctk.CTkFrame(
            self, 
            fg_color=COULEURS["fond_principal"]
        )
        self.cadre_saisie.pack(fill='x', padx=10, pady=10)
        
        self.saisie_message = ctk.CTkEntry(
            self.cadre_saisie, 
            placeholder_text="Saisissez votre commande ou message...",
            width=1100,
            fg_color=COULEURS["fond_chat"],
            text_color=COULEURS["texte_principal"],
            placeholder_text_color=COULEURS["texte_secondaire"]
        )
        self.saisie_message.pack(side='left', padx=10, expand=True, fill='x')
        self.saisie_message.bind('<Return>', self.envoyer_message)
        self.saisie_message.bind('<Tab>', self._autocompletion)
        self.saisie_message.bind('<KeyRelease>', self._suggestion_commande)
        
        self.bouton_envoi = ctk.CTkButton(
            self.cadre_saisie, 
            text="Envoyer", 
            command=self.envoyer_message,
            fg_color=COULEURS["accent_secondaire"],
            hover_color=COULEURS["accent_primaire"],
            text_color=COULEURS["texte_principal"]
        )
        self.bouton_envoi.pack(side='right', padx=10)
        
        # Liste de suggestions
        self.liste_suggestions = tk.Listbox(
            self, 
            height=5, 
            width=50, 
            bg=COULEURS["fond_chat"], 
            fg=COULEURS["texte_principal"],
            selectbackground=COULEURS["accent_primaire"]
        )
        self.liste_suggestions.pack(fill='x', padx=10, pady=5)
        self.liste_suggestions.bind('<Double-1>', self._selectionner_suggestion)
        
        # Barre d'outils pour les contextes
        self.barre_outils = ctk.CTkFrame(
            self, 
            fg_color="transparent", 
            corner_radius=0
        )
        self.barre_outils.pack(fill='x', padx=10, pady=5)
        
        # Sélecteur de contexte
        ctk.CTkLabel(
            self.barre_outils, 
            text="Contexte", 
            text_color=COULEURS["texte_principal"],
            font=("Arial", 12, "bold")
        ).pack(side='left', padx=(10,5))
        
        self.selection_contexte = ctk.CTkComboBox(
            self.barre_outils, 
            values=["Développement", "Cybersécurité", "Data Science", "Cloud"],
            command=self._changer_contexte,
            width=250,
            state="readonly"
        )
        self.selection_contexte.pack(side='left', padx=5)
        
        # Bouton pour charger un fichier
        self.bouton_charger = ctk.CTkButton(
            self.barre_outils, 
            text=" Charger Fichier", 
            command=self._charger_fichier,
            width=150
        )
        self.bouton_charger.pack(side='left', padx=5)
        
        # Bouton pour résumer le fichier
        self.bouton_resumer = ctk.CTkButton(
            self.barre_outils, 
            text=" Résumer", 
            command=self._resumer_fichier,
            width=150,
            state='disabled'
        )
        self.bouton_resumer.pack(side='left', padx=5)
        
        # Bouton pour générer des images
        self.bouton_generer_image = ctk.CTkButton(
            self.barre_outils, 
            text=" Générer Image", 
            command=self._afficher_dialogue_generation_image,
            width=150
        )
        self.bouton_generer_image.pack(side='left', padx=5)
        
        # Bouton pour le mode écriture
        self.bouton_ecriture = ctk.CTkButton(
            self.barre_outils,
            text="Mode Écriture", 
            command=self._ajouter_mode_ecriture,
            width=120
        )
        self.bouton_ecriture.pack(side='left', padx=5)
        
        # Case à cocher pour la recherche web
        self.var_recherche_web = tk.BooleanVar(value=False)
        self.case_recherche_web = ctk.CTkCheckBox(
            self.barre_outils, 
            text="Recherche Web", 
            variable=self.var_recherche_web,
            onvalue=True, 
            offvalue=False,
            width=120
        )
        self.case_recherche_web.pack(side='left', padx=5)
        
        # Frame pour afficher les images
        self.cadre_images = ctk.CTkScrollableFrame(
            self, 
            fg_color="transparent", 
            orientation="horizontal"
        )
        self.cadre_images.pack(fill='x', padx=10, pady=5)
        
        # File d'attente pour la génération d'image
        self.file_generation_image = queue.Queue()
        
        # Verrou pour éviter les générations multiples
        self.generation_en_cours = False
        
    def _ajouter_message(self, expediteur, message):
        cadre_message = ctk.CTkFrame(
            self.cadre_chat, 
            fg_color=COULEURS["fond_principal"]
        )
        cadre_message.pack(fill='x', padx=5, pady=5)

        libelle_expediteur = ctk.CTkLabel(
            cadre_message, 
            text=f"{expediteur} :", 
            font=('Helvetica', 12, 'bold'),
            text_color=COULEURS["accent_secondaire"]
        )
        libelle_expediteur.pack(anchor='w')

        libelle_message = ctk.CTkLabel(
            cadre_message, 
            text=message, 
            wraplength=1200, 
            justify='left',
            text_color=COULEURS["texte_principal"]
        )
        libelle_message.pack(anchor='w')
        
    def envoyer_message(self, event=None):
        message = self.saisie_message.get()
        
        # Vérifier si la recherche web est activée
        if self.var_recherche_web.get():
            self._effectuer_recherche_web(message)
            return
        
        if message:
            # Vérifier si c'est une commande spéciale
            resultat_commande = self._traiter_commande_speciale(message)
            
            if resultat_commande:
                self._ajouter_message("Assistant", resultat_commande)
            else:
                # Ajouter le message de l'utilisateur
                self._ajouter_message("Utilisateur", message)
                
                # Désactiver la saisie pendant le traitement
                self.saisie_message.configure(state='disabled')
                
                # Lancer la génération de réponse dans un thread séparé
                threading.Thread(target=self._generer_reponse_ia, args=(message,), daemon=True).start()
            
            # Effacer la saisie
            self.saisie_message.delete(0, tk.END)
        
        if message and self.fichier_charge:
            # Si un fichier est chargé, proposer des interactions
            interactions = [
                "Peux-tu résumer ce document ?",
                "Quels sont les points clés ?",
                "Explique-moi ce document"
            ]
            
            # Ajouter des suggestions d'interaction avec le document
            self.liste_suggestions.delete(0, tk.END)
            for interaction in interactions:
                self.liste_suggestions.insert(tk.END, interaction)
        
        # Vérifier si le message est une demande de génération d'image
        mots_cles_image = ["génère une image", "crée une image", "dessine"]
        if any(mot in message.lower() for mot in mots_cles_image):
            # Extraire la description de l'image
            description = message.lower().replace("génère une image", "").replace("crée une image", "").replace("dessine", "").strip()
            
            # Ouvrir le dialogue de génération avec la description
            self._afficher_dialogue_generation_image()
            self.description_image.delete(0, tk.END)
            self.description_image.insert(0, description)
            
            return
        
    def _generer_reponse_ia(self, message):
        try:
            # Récupérer la configuration
            config = self.gestionnaire_config.obtenir_configuration()
            modele = config.get('modeles', {}).get('actif', 'llama3.2')

            # Prompts spécifiques par contexte
            prompts_contextuels = {
                "Développement": """
                Tu es un expert en développement logiciel. 
                Réponds aux questions techniques avec précision et des exemples de code.
                Utilise les bonnes pratiques de programmation et explique clairement.
                """,
                "Cybersécurité": """
                Tu es un expert en cybersécurité. 
                Analyse les questions sous l'angle de la sécurité informatique.
                Fournis des recommandations de sécurité et explique les risques potentiels.
                """,
                "Data Science": """
                Tu es un data scientist expérimenté. 
                Réponds aux questions avec des insights analytiques et statistiques.
                Propose des approches méthodologiques et des techniques d'analyse.
                """,
                "Cloud": """
                Tu es un architecte cloud et expert en infrastructure.
                Réponds aux questions sur les technologies cloud, l'architecture et le déploiement.
                Donne des conseils pratiques sur les solutions cloud.
                """
            }

            # Prompt de système pour le contexte actuel
            prompt_systeme = prompts_contextuels.get(self.contexte_actuel, """
            Tu es un assistant IA conversationnel. 
            Réponds TOUJOURS en français avec précision et clarté.
            """)

            # Log de débogage
            print(f"[DEBUG] Modèle utilisé : {modele}")
            print(f"[DEBUG] Contexte : {self.contexte_actuel}")
            print(f"[DEBUG] Message : {message}")

            # Générer la réponse avec Ollama
            try:
                reponse = self.gestionnaire_ollama.chat(
                    model=modele,
                    messages=[
                        {'role': 'system', 'content': prompt_systeme},
                        {'role': 'user', 'content': message}
                    ]
                )
                
                # Log de débogage supplémentaire
                print(f"[DEBUG] Type de réponse : {type(reponse)}")
                print(f"[DEBUG] Contenu de la réponse : {reponse}")

                # Vérification robuste de la réponse
                if isinstance(reponse, dict) and 'message' in reponse:
                    contenu_reponse = reponse['message'].get('content', 'Réponse vide')
                    self._ajouter_message("Assistant IA", contenu_reponse)
                elif hasattr(reponse, 'message'):
                    self._ajouter_message("Assistant IA", reponse.message)
                else:
                    self._ajouter_message("Assistant IA", str(reponse))

            except AttributeError as attr_err:
                print(f"[ERREUR] Erreur d'attribut : {attr_err}")
                self._ajouter_message("Assistant IA", f"Erreur d'attribut : {attr_err}")
            except TypeError as type_err:
                print(f"[ERREUR] Erreur de type : {type_err}")
                self._ajouter_message("Assistant IA", f"Erreur de type : {type_err}")

        except Exception as e:
            import traceback
            print(f"[ERREUR GLOBALE] {str(e)}")
            traceback.print_exc()
            self._ajouter_message("Assistant IA", f"Erreur : {str(e)}")
        finally:
            # Réactiver la saisie
            self.saisie_message.configure(state='normal')

    def _changer_contexte(self, nouveau_contexte):
        """Change le contexte de la conversation"""
        self.contexte_actuel = nouveau_contexte
        self._ajouter_message("Système", f"Contexte changé pour : {nouveau_contexte}")
        
    def _suggestion_commande(self, event):
        """Génère des suggestions de commandes basées sur le contexte"""
        commande = self.saisie_message.get()
        suggestions = self._obtenir_suggestions(commande)
        
        self.liste_suggestions.delete(0, tk.END)
        for suggestion in suggestions:
            self.liste_suggestions.insert(tk.END, suggestion)
        
    def _obtenir_suggestions(self, commande):
        """Retourne les suggestions de commandes par contexte"""
        suggestions = {
            "Développement": [
                "/generer_code_python", 
                "/debugger", 
                "/generer_classe", 
                "/refactoriser"
            ],
            "Cybersécurité": [
                "/scanner_vulnerabilite", 
                "/generer_politique", 
                "/analyser_logs"
            ],
            "Data Science": [
                "/analyser_donnees", 
                "/generer_graphique", 
                "/entrainer_modele"
            ],
            "Cloud": [
                "/deployer_container", 
                "/configurer_kubernetes", 
                "/auditer_securite"
            ]
        }
        return [cmd for cmd in suggestions.get(self.contexte_actuel, []) if cmd.startswith(commande)]
    
    def _selectionner_suggestion(self, event):
        """Sélectionne une suggestion de la liste"""
        selection = self.liste_suggestions.get(self.liste_suggestions.curselection())
        self.saisie_message.delete(0, tk.END)
        self.saisie_message.insert(0, selection)
    
    def _autocompletion(self, event):
        """Complète automatiquement la commande"""
        commande = self.saisie_message.get()
        suggestions = self._obtenir_suggestions(commande)
        
        if suggestions:
            self.saisie_message.delete(0, tk.END)
            self.saisie_message.insert(0, suggestions[0])
        
        return 'break'
    
    def _traiter_commande_speciale(self, commande):
        """Traite les commandes spéciales"""
        commandes_speciales = {
            "/generer_code_python": self._generer_code_python,
            "/debugger": self._mode_debug,
            "/generer_classe": self._generer_classe,
            "/scanner_vulnerabilite": self._scanner_vulnerabilite,
            "/analyser_donnees": self._analyser_donnees
        }
        
        for prefixe, fonction in commandes_speciales.items():
            if commande.startswith(prefixe):
                return fonction(commande[len(prefixe):].strip())
        
        return None
    
    def _generer_code_python(self, contexte):
        """Génère un exemple de code Python"""
        exemples = {
            "": "class ExempleClasse:\n    def __init__(self):\n        pass\n\n    def methode_exemple(self):\n        return 'Hello, World!'",
            "api": "import flask\n\napp = flask.Flask(__name__)\n\n@app.route('/')\ndef hello_world():\n    return 'Hello, World!'",
            "machine_learning": "import numpy as np\nfrom sklearn.linear_model import LinearRegression\n\n# Exemple de régression linéaire\nX = np.array([[1], [2], [3], [4]])\ny = np.array([2, 4, 5, 4])\n\nmodel = LinearRegression()\nmodel.fit(X, y)"
        }
        code = exemples.get(contexte, exemples[""])
        return f"```python\n{code}\n```"

    def _mode_debug(self, code):
        """Simule un débogage"""
        erreurs_simulees = [
            "Aucune erreur détectée.",
            "Attention : Variable non initialisée.",
            "Syntaxe incorrecte à la ligne 3.",
            "Importation manquante : numpy"
        ]
        return random.choice(erreurs_simulees)

    def _generer_classe(self, nom_classe):
        """Génère une classe Python"""
        nom_classe = nom_classe or "MaClasse"
        code = f"""
class {nom_classe}:
    def __init__(self):
        \"\"\"Constructeur de la classe {nom_classe}\"\"\"
        pass

    def methode_exemple(self):
        \"\"\"Méthode exemple\"\"\"
        return "Méthode de {nom_classe}"
"""
        return f"```python\n{code}\n```"

    def _scanner_vulnerabilite(self, cible):
        """Simule un scan de vulnérabilité"""
        resultats = [
            "Aucune vulnérabilité critique détectée.",
            "Attention : Potentielle faille XSS détectée.",
            "Risque de injection SQL identifié.",
            "Configuration de sécurité recommandée."
        ]
        return random.choice(resultats)

    def _analyser_donnees(self, description):
        """Simule une analyse de données"""
        analyses = [
            "Distribution normale détectée.",
            "Corrélation significative entre variables.",
            "Données nécessitant un prétraitement.",
            "Recommandation : Réduction de dimensionnalité"
        ]
        return random.choice(analyses)

    def _charger_fichier(self):
        """Charge un fichier et prépare pour résumé"""
        try:
            print(" Ouverture du sélecteur de fichier")
            chemin_fichier = filedialog.askopenfilename(
                title="Sélectionner un fichier",
                filetypes=[
                    ("Tous les fichiers", "*.*"),
                    ("Fichiers PDF", "*.pdf"),
                    ("Documents Word", "*.docx"),
                    ("Fichiers Texte", "*.txt"),
                ]
            )
            
            print(f" Chemin du fichier sélectionné : {chemin_fichier}")
            
            if not chemin_fichier:
                print(" Aucun fichier sélectionné")
                return
            
            # Extraire le texte selon le type de fichier
            print(" Début de l'extraction de texte")
            self.fichier_charge = self._extraire_texte(chemin_fichier)
            print(f" Longueur du texte extrait : {len(self.fichier_charge) if self.fichier_charge else 0}")
            
            # Activer le bouton de résumé
            print(" Activation du bouton résumer")
            self.bouton_resumer.configure(state='normal')
            
            # Afficher un message de confirmation
            print(f" Fichier chargé : {os.path.basename(chemin_fichier)}")
            self._ajouter_message("Système", f"Fichier chargé : {os.path.basename(chemin_fichier)}")
            
            # Mettre à jour le dernier chemin de fichier
            self.dernier_chemin_fichier = chemin_fichier
            
        except Exception as e:
            print(f" ERREUR CRITIQUE DE CHARGEMENT : {e}")
            print(f"Type d'erreur : {type(e)}")
            import traceback
            traceback.print_exc()  # Affiche la trace complète de l'erreur
            self._ajouter_message("Système", f"Erreur de chargement : {str(e)}")

    def _extraire_texte(self, chemin_fichier):
        """
        Extrait le texte d'un fichier en fonction de son extension.
        
        Args:
            chemin_fichier (str): Chemin complet du fichier à extraire
        
        Returns:
            str: Texte extrait du fichier
        """
        print(f" Début de l'extraction pour : {chemin_fichier}")
        
        # Obtenir l'extension du fichier
        extension = os.path.splitext(chemin_fichier)[1].lower()
        print(f" Extension du fichier : {extension}")
        
        try:
            texte_extrait = ""
            
            # Extraction PDF avec pdfplumber
            if extension == '.pdf':
                print(" Extraction PDF avec pdfplumber")
                import pdfplumber
                
                with pdfplumber.open(chemin_fichier) as pdf:
                    # Extraire le texte de chaque page
                    pages_texte = []
                    for page in pdf.pages:
                        # Essayer plusieurs méthodes d'extraction
                        page_texte = page.extract_text()
                        
                        # Si extract_text() échoue, essayer extract_words()
                        if not page_texte or not page_texte.strip():
                            words = page.extract_words()
                            page_texte = ' '.join([word['text'] for word in words])
                        
                        if page_texte and page_texte.strip():
                            pages_texte.append(page_texte)
                    
                    print(f" Pages PDF extraites : {len(pages_texte)}")
                    
                    # Concaténer les pages
                    texte_extrait = "\n".join(pages_texte)
                    
                    # Nettoyer et formater le texte
                    texte_extrait = texte_extrait.replace('\n', ' ').strip()
            
            # Extraction DOCX
            elif extension == '.docx':
                print(" Extraction DOCX avec python-docx")
                import docx
                doc = docx.Document(chemin_fichier)
                texte_extrait = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            
            # Extraction TXT et MD
            elif extension in ['.txt', '.md']:
                print(f" Extraction {extension}")
                with open(chemin_fichier, 'r', encoding='utf-8') as fichier:
                    texte_extrait = fichier.read().strip()
            
            else:
                print(f" Type de fichier non supporté : {extension}")
                return ""
            
            # Vérifier la longueur du texte extrait
            print(f" Longueur du texte extrait : {len(texte_extrait)}")
            
            # Nettoyer le texte
            texte_extrait = re.sub(r'\s+', ' ', texte_extrait).strip()
            
            # Vérifier si le texte extrait est significatif
            if len(texte_extrait) < 50:
                print(" Texte extrait trop court, possible erreur d'extraction")
                return ""
            
            return texte_extrait
        
        except Exception as e:
            print(f" Erreur lors de l'extraction : {e}")
            import traceback
            traceback.print_exc()
            return ""

    def _lire_fichier_avec_encodage(self, chemin_fichier):
        """
        Lire un fichier avec différents encodages
        
        :param chemin_fichier: Chemin complet du fichier
        :return: Contenu du fichier
        """
        # Liste des encodages à essayer
        encodages = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        # Vérifier si le fichier existe
        if not os.path.exists(chemin_fichier):
            raise FileNotFoundError(f"Le fichier {chemin_fichier} n'existe pas")
        
        # Essayer différents encodages
        for encodage in encodages:
            try:
                with open(chemin_fichier, 'r', encoding=encodage) as fichier:
                    return fichier.read()
            except UnicodeDecodeError:
                # Passer à l'encodage suivant si celui-ci échoue
                continue
        
        # Si aucun encodage ne fonctionne
        raise UnicodeDecodeError(f"Impossible de lire le fichier avec les encodages : {encodages}")
    
    def _nettoyer_et_corriger_chemin(self, chemin_original):
        """
        Nettoie et corrige les chemins problématiques
        
        :param chemin_original: Chemin original potentiellement incorrect
        :return: Chemin corrigé
        """
        # Débogage initial
        print(" Début de la correction de chemin")
        print(f" Chemin original brut : '{repr(chemin_original)}'")
        
        # Vérifier si l'entrée est None ou vide
        if not chemin_original:
            print(" Chemin d'entrée vide ou None")
            return None
        
        # Convertir en chaîne si ce n'est pas déjà le cas
        try:
            chemin_original = str(chemin_original)
        except Exception as e:
            print(f" Erreur de conversion en chaîne : {e}")
            return None
        
        # Supprimer les espaces supplémentaires au début et à la fin
        chemin_corrige = chemin_original.strip()
        print(f" Après strip : '{chemin_corrige}'")
        
        # Correction des séparateurs
        chemin_corrige = chemin_corrige.replace('\\', os.path.sep)
        chemin_corrige = chemin_corrige.replace('/', os.path.sep)
        print(f" Après correction des séparateurs : '{chemin_corrige}'")
        
        # Correction des espaces dans le chemin
        chemin_corrige = re.sub(r'\s+', ' ', chemin_corrige)  # Remplace plusieurs espaces par un seul
        print(f" Après correction des espaces : '{chemin_corrige}'")
        
        # Correction spécifique pour les chemins Windows
        if chemin_corrige.startswith('C:') and not chemin_corrige.startswith('C:\\'):
            chemin_corrige = chemin_corrige.replace('C:', 'C:\\')
        print(f" Après correction Windows : '{chemin_corrige}'")
        
        # Corrections de chemins courants
        corrections_chemins = {
            'Users\\gabrilOneDrive': 'Users\\gabri\\OneDrive',
            'OneDrive \\Bureau': 'OneDrive\\Bureau',
            'OneDrive \Bureau': 'OneDrive\\Bureau',
            'Cinquiem e semaine': 'Cinquième semaine',
            'miniprojet ': 'miniprojet'
        }
        
        for motif, correction in corrections_chemins.items():
            if motif in chemin_corrige:
                chemin_corrige = chemin_corrige.replace(motif, correction)
                print(f" Correction '{motif}' -> '{correction}'")
        
        # Normaliser le chemin
        try:
            chemin_corrige = os.path.normpath(chemin_corrige)
            print(f" Après normalisation : '{chemin_corrige}'")
        except Exception as e:
            print(f" Erreur de normalisation : {e}")
        
        # Vérification de l'existence du fichier/répertoire
        if os.path.exists(chemin_corrige):
            print(f" Chemin validé : {chemin_corrige}")
            return chemin_corrige
        
        # Tentative de correction avancée
        repertoire_parent = os.path.dirname(chemin_corrige)
        print(f" Répertoire parent : '{repertoire_parent}'")
        
        try:
            if os.path.exists(repertoire_parent):
                # Lister les fichiers du répertoire
                fichiers_repertoire = os.listdir(repertoire_parent)
                print(f" Fichiers du répertoire parent ({repertoire_parent}) :")
                for fichier in fichiers_repertoire:
                    print(f"- {fichier}")
                
                # Recherche de fichiers similaires
                nom_fichier = os.path.basename(chemin_corrige)
                print(f" Recherche de fichiers similaires à : '{nom_fichier}'")
                
                for fichier in fichiers_repertoire:
                    # Comparaison insensible à la casse et aux espaces
                    if (fichier.lower().replace(' ', '') == 
                        nom_fichier.lower().replace(' ', '')):
                        chemin_corrige = os.path.join(repertoire_parent, fichier)
                        print(f" Fichier similaire trouvé : {chemin_corrige}")
                        return chemin_corrige
        except Exception as e:
            print(f" Erreur lors de la recherche : {e}")
        
        # Dernier recours : message d'erreur
        print(" Impossible de trouver le fichier")
        return None
    
    def _verifier_origine_fichier(self, chemin_fichier):
        """
        Vérification approfondie de l'origine du fichier
        
        :param chemin_fichier: Chemin du fichier à vérifier
        :return: Informations détaillées sur le fichier
        """
        # Vérifier si le chemin est None ou vide
        if not chemin_fichier:
            print(" Aucun chemin de fichier fourni")
            raise FileNotFoundError("Aucun chemin de fichier valide")
        
        # Nettoyer et corriger le chemin
        chemin_corrige = self._nettoyer_et_corriger_chemin(chemin_fichier)
        
        # Vérifier à nouveau si le chemin est None après correction
        if chemin_corrige is None:
            print(" Impossible de corriger le chemin du fichier")
            raise FileNotFoundError(f"Chemin incorrect : {chemin_fichier}")
        
        # Collecter toutes les informations possibles
        infos_fichier = {
            "chemin_original": chemin_fichier,
            "chemin_normalise": os.path.normpath(chemin_corrige),
            "chemin_absolu": os.path.abspath(chemin_corrige),
            "repertoire_courant": os.getcwd(),
            "repertoire_parent": os.path.dirname(chemin_corrige),
            "fichier_existe": os.path.exists(chemin_corrige),
            "est_fichier": os.path.isfile(chemin_corrige),
            "est_lisible": os.access(chemin_corrige, os.R_OK) if os.path.exists(chemin_corrige) else False
        }
        
        # Ajouter des messages de débogage détaillés
        messages_debug = [
            f" Origine du fichier : {infos_fichier['chemin_original']}",
            f" Chemin normalisé : {infos_fichier['chemin_normalise']}",
            f" Chemin absolu : {infos_fichier['chemin_absolu']}",
            f" Répertoire courant : {infos_fichier['repertoire_courant']}",
            f" Répertoire parent : {infos_fichier['repertoire_parent']}",
            f" Fichier existe : {infos_fichier['fichier_existe']}",
            f" Est un fichier : {infos_fichier['est_fichier']}",
            f" Fichier lisible : {infos_fichier['est_lisible']}"
        ]
        
        # Lister les fichiers du répertoire parent
        try:
            repertoire_parent = infos_fichier['repertoire_parent']
            if repertoire_parent and os.path.exists(repertoire_parent):
                fichiers_repertoire = os.listdir(repertoire_parent)
                print(f" Fichiers du répertoire : {fichiers_repertoire}")
        except Exception as e:
            messages_debug.append(f" Erreur de listage : {e}")
        
        # Ajouter tous les messages de débogage
        for message in messages_debug:
            print(message)
        
        # Vérification finale
        if not infos_fichier['est_fichier'] or not infos_fichier['est_lisible']:
            raise FileNotFoundError(
                f"Impossible de lire le fichier {infos_fichier['chemin_absolu']}. "
                "Vérifiez le chemin, les permissions et l'existence du fichier."
            )
        
        return infos_fichier
    
    def _resumer_fichier(self, chemin_fichier=None):
        """
        Résume le contenu du fichier chargé
        
        Args:
            chemin_fichier (str, optional): Chemin du fichier à résumer. 
            Si None, tentera de retrouver le fichier chargé.
        """
        # Vérifier si un chemin de fichier est fourni
        if chemin_fichier is None:
            # Liste des chemins possibles
            derniers_chemins = [
                chemin for chemin in [
                    getattr(self, 'dernier_chemin_fichier', None),
                    getattr(self, 'chemin_fichier_charge', None)
                ] if chemin is not None
            ]
            
            # Vérifier si un fichier est chargé
            if not derniers_chemins:
                messagebox.showerror("Erreur", "Veuillez d'abord charger un fichier")
                return "Aucun fichier chargé"
            
            # Prendre le premier chemin disponible
            chemin_fichier = derniers_chemins[0]
        
        print(f" Chemin du fichier à résumer : {chemin_fichier}")
        
        try:
            # Extraire le texte du fichier
            contenu = self._extraire_texte(chemin_fichier)
            
            # Vérifier si le texte a été extrait avec succès
            if not contenu or len(contenu) < 50:
                self._ajouter_message("Système", "Impossible d'extraire le texte du fichier.")
                return "Aucun contenu significatif n'a pu être extrait du fichier."
            
            # Tronquer le texte si trop long (limiter à 10000 caractères)
            if len(contenu) > 10000:
                contenu = contenu[:10000]
            
            # Préparer le prompt de résumé
            prompt_systeme = f"""
            Tu es un assistant professionnel spécialisé dans la rédaction de résumés en français.

            OBJECTIFS PRINCIPAUX :
            - Générer un résumé précis et informatif du document
            - Capturer l'essence et les points clés du texte
            - Produire un résumé en français clair et structuré

            INSTRUCTIONS DÉTAILLÉES :
            1. Analyse approfondie du document
            2. Identifie le type et le contexte du document
            3. Extrais les informations essentielles
            4. Rédige un résumé professionnel qui répond aux questions :
               - Quel est le sujet principal ?
               - Quels sont les points clés ?
               - Quelles sont les conclusions ou implications ?

            RÈGLES DE RÉDACTION :
            - Longueur : 200-400 mots
            - Langage : Français professionnel
            - Structure : Introduction, points principaux, conclusion
            - Objectivité : Neutre et factuel
            - Clarté : Compréhensible par un public général

            CONTEXTE DU DOCUMENT :
            - Nom : {os.path.basename(chemin_fichier)}
            - Type de fichier : {os.path.splitext(chemin_fichier)[1]}
            - Longueur : {len(contenu)} caractères

            CONSIGNE FINALE :
            Résume le document de manière à ce qu'une personne puisse comprendre son contenu essentiel en quelques minutes.

            CONTENU DU DOCUMENT :
            {contenu}
            """
            
            # Générer le résumé avec Ollama
            reponse = self._generer_reponse_ollama(
                prompt_systeme, 
                modele="mistral", 
                temperature=0.3
            )
            
            # Ajouter le résumé au chat
            self._ajouter_message("Système", f"📄 Résumé du fichier {os.path.basename(chemin_fichier)} :\n\n{reponse}")
            
            return reponse
        
        except Exception as e:
            print(f" Erreur lors du résumé : {e}")
            import traceback
            traceback.print_exc()
            message_erreur = f"Erreur de résumé : {str(e)}"
            self._ajouter_message("Système", message_erreur)
            messagebox.showerror("Erreur de Résumé", message_erreur)
            return message_erreur
    
    def _afficher_dialogue_generation_image(self):
        """Ouvre un dialogue pour générer une image"""
        self.dialogue_image = ctk.CTkToplevel(self)
        self.dialogue_image.title("Génération d'Image")
        self.dialogue_image.geometry("500x350")
        
        # Titre
        ctk.CTkLabel(
            self.dialogue_image, 
            text="Générer une Image", 
            font=("Arial", 16, "bold")
        ).pack(pady=(20,10))
        
        # Champ de saisie pour la description
        self.description_image = ctk.CTkEntry(
            self.dialogue_image, 
            placeholder_text="Décrivez l'image que vous voulez générer...",
            width=400
        )
        self.description_image.pack(pady=10)
        
        # Options de style
        self.style_image = ctk.CTkComboBox(
            self.dialogue_image,
            values=[
                "Réaliste", 
                "Cartoon", 
                "Pixel Art", 
                "Aquarelle", 
                "Dessin au trait"
            ],
            width=400
        )
        self.style_image.pack(pady=10)
        
        # Barre de progression
        self.barre_progression_image = ctk.CTkProgressBar(
            self.dialogue_image, 
            width=400, 
            mode='determinate'
        )
        self.barre_progression_image.pack(pady=10)
        self.barre_progression_image.set(0)
        
        # Étiquette de progression
        self.etiquette_progression = ctk.CTkLabel(
            self.dialogue_image, 
            text="En attente de génération...", 
            font=("Arial", 10)
        )
        self.etiquette_progression.pack(pady=5)
        
        # Bouton de génération
        self.bouton_generer = ctk.CTkButton(
            self.dialogue_image, 
            text="Générer", 
            command=self._preparer_generation_image,
            width=400
        )
        self.bouton_generer.pack(pady=10)
        
    def _preparer_generation_image(self):
        """Prépare la génération d'image avec progression"""
        # Vérifier si une génération est déjà en cours
        if self.generation_en_cours:
            print(" Une génération d'image est déjà en cours.")
            return
        
        description = self.description_image.get()
        style = self.style_image.get()
        
        if not description:
            print(" Veuillez entrer une description.")
            return
        
        # Marquer la génération comme en cours
        self.generation_en_cours = True
        
        # Désactiver le bouton pendant la génération
        self.bouton_generer.configure(state='disabled')
        
        # Réinitialiser la barre de progression
        self.barre_progression_image.set(0)
        
        # Lancer la génération dans un thread séparé
        thread_generation = threading.Thread(
            target=self._generer_image_avec_progression, 
            args=(description, style)
        )
        thread_generation.start()
        
        # Mettre à jour la progression
        self._mettre_a_jour_progression()
        
    def _mettre_a_jour_progression(self):
        """Met à jour la barre de progression"""
        try:
            progression = self.file_generation_image.get_nowait()
            
            if progression == 100:
                # Génération terminée
                self.barre_progression_image.set(1)
                self.etiquette_progression.configure(text="Génération terminée !")
                
                # Fermer automatiquement la fenêtre après un court délai
                self.dialogue_image.after(1500, self._finaliser_generation)
                return
            
            # Mettre à jour la progression
            self.barre_progression_image.set(progression / 100)
            self.etiquette_progression.configure(text=f"Génération en cours... {progression}%")
            
        except queue.Empty:
            pass
        
        # Programmer la prochaine mise à jour
        self.dialogue_image.after(100, self._mettre_a_jour_progression)
        
    def _finaliser_generation(self):
        """Finalise la génération et réinitialise les états"""
        # Fermer la fenêtre
        self.dialogue_image.destroy()
        
        # Réactiver la possibilité de générer une nouvelle image
        self.generation_en_cours = False
        
    def _generer_image_avec_progression(self, description, style):
        """Génère une image avec progression simulée"""
        try:
            # Simulation de progression
            for i in range(0, 101, 10):
                self.file_generation_image.put(i)
                time.sleep(0.5)  # Simuler un temps de traitement
            
            # API Hugging Face pour génération d'image
            API_URL = "https://api-inference.huggingface.co/models/stabilityai/stable-diffusion-xl-base-1.0"
            headers = {
                "Authorization": "Bearer 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",
                "Content-Type": "application/json"
            }
            
            payload = {
                "inputs": f"{description} dans un style {style}, high quality, detailed"
            }
            
            # Envoyer la requête
            reponse = requests.post(API_URL, headers=headers, json=payload)
            
            if reponse.status_code != 200:
                print(f" Erreur API : {reponse.text}")
                return
            
            # Convertir l'image
            image = Image.open(BytesIO(reponse.content))
            
            # Sauvegarder l'image
            dossier_images = os.path.join(os.getcwd(), "images_generees")
            os.makedirs(dossier_images, exist_ok=True)
            chemin_image = os.path.join(dossier_images, f"image_{len(os.listdir(dossier_images)) + 1}.png")
            image.save(chemin_image)
            
            # Fermer la fenêtre modale
            if hasattr(self, 'dialogue_image') and self.dialogue_image.winfo_exists():
                self.dialogue_image.destroy()
            
            # Créer un cadre de message pour l'image
            cadre_message = ctk.CTkFrame(
                self.cadre_chat, 
                fg_color=COULEURS["fond_chat"],
                corner_radius=10
            )
            cadre_message.pack(fill='x', padx=10, pady=5, anchor='w')
            
            # Étiquette pour l'expéditeur
            ctk.CTkLabel(
                cadre_message, 
                text="Générateur d'Image", 
                font=("Arial", 10, "bold"),
                text_color=COULEURS["texte_secondaire"]
            ).pack(anchor='w', padx=10, pady=(5,0))
            
            # Afficher l'image dans le chat
            image_tk = ctk.CTkImage(
                Image.open(chemin_image), 
                size=(300, 300)
            )
            
            label_image = ctk.CTkLabel(
                cadre_message, 
                image=image_tk, 
                text=""
            )
            label_image.pack(padx=10, pady=5)
            
            # Description de l'image
            ctk.CTkLabel(
                cadre_message, 
                text=f"Description : {description}\nStyle : {style}", 
                font=("Arial", 10),
                text_color=COULEURS["texte_secondaire"]
            ).pack(anchor='w', padx=10, pady=(0,5))
            
            # Faire défiler jusqu'en bas
            self.cadre_chat.update()
            self.cadre_chat._parent_canvas.yview_moveto(1.0)
            
        except Exception as e:
            print(f" Erreur de génération d'image : {str(e)}")
            
        finally:
            # S'assurer que la progression atteint 100%
            self.file_generation_image.put(100)
            
            # Réinitialiser le verrou de génération
            self.generation_en_cours = False
            
    def _effectuer_recherche_web(self, requete):
        """Effectue une recherche web et affiche les résultats"""
        try:
            # Clé API SerpAPI (vous devrez la remplacer)
            api_key = "db9e53e5ebb0deb3d3d6c45ab14e5822b5e0e6b0862771f7ebf723c49ee1855f"  # À remplacer par votre clé
            
            # Paramètres de recherche
            params = {
                "engine": "google",
                "q": requete,
                "api_key": api_key,
                "num": 5  # Limiter à 5 résultats
            }
            
            # URL de l'API SerpAPI
            url = "https://serpapi.com/search"
            
            # Effectuer la requête
            reponse = requests.get(url, params=params)
            
            if reponse.status_code != 200:
                print(f" Erreur de recherche : {reponse.text}")
                return
            
            # Convertir la réponse en JSON
            resultats_json = reponse.json()
            
            # Vérifier si des résultats sont présents
            if 'organic_results' not in resultats_json:
                print(" Aucun résultat trouvé.")
                return
            
            # Créer un cadre de message pour les résultats
            cadre_resultats = ctk.CTkFrame(
                self.cadre_chat, 
                fg_color=COULEURS["fond_chat"],
                corner_radius=10
            )
            cadre_resultats.pack(fill='x', padx=10, pady=5, anchor='w')
            
            # Étiquette pour l'expéditeur
            ctk.CTkLabel(
                cadre_resultats, 
                text="Résultats de Recherche Web", 
                font=("Arial", 10, "bold"),
                text_color=COULEURS["texte_secondaire"]
            ).pack(anchor='w', padx=10, pady=(5,0))
            
            # Parcourir les résultats
            for i, resultat in enumerate(resultats_json['organic_results'][:5], 1):
                # Cadre pour chaque résultat
                cadre_resultat = ctk.CTkFrame(
                    cadre_resultats, 
                    fg_color="transparent"
                )
                cadre_resultat.pack(fill='x', padx=10, pady=2, anchor='w')
                
                # Titre du résultat (cliquable)
                lien_titre = ctk.CTkButton(
                    cadre_resultat,
                    text=f"{i}. {resultat.get('title', 'Pas de titre')[:100]}...", 
                    command=lambda url=resultat.get('link', '#'): webbrowser.open(url),
                    fg_color="transparent",
                    hover_color=COULEURS["accent_secondaire"],
                    text_color=COULEURS["accent_primaire"],
                    anchor='w',
                    width=500,
                    height=30
                )
                lien_titre.pack(side='left', padx=(0,5))
                
                # Description courte
                ctk.CTkLabel(
                    cadre_resultat, 
                    text=resultat.get('snippet', 'Pas de description')[:200] + "...", 
                    font=("Arial", 10),
                    text_color=COULEURS["texte_secondaire"],
                    wraplength=500,
                    width=500
                ).pack(side='left')
            
            # Faire défiler jusqu'en bas
            self.cadre_chat.update()
            self.cadre_chat._parent_canvas.yview_moveto(1.0)
            
        except Exception as e:
            print(f" Erreur de recherche web : {str(e)}")
            # Ajouter un message de débogage détaillé
            import traceback
            traceback.print_exc()
            
    def _rechercher_web(self, requete):
        """Effectue une recherche web et retourne les résultats"""
        try:
            # Encoder la requête pour l'URL
            requete_encodee = urllib.parse.quote(requete)
            
            # URL de recherche Google
            url_recherche = f"https://www.google.com/search?q={requete_encodee}"
            
            # En-têtes pour simuler un navigateur
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'Accept-Language': 'fr-FR,fr;q=0.9,en-US;q=0.8,en;q=0.7'
            }
            
            # Effectuer la requête
            reponse = requests.get(url_recherche, headers=headers)
            
            # Vérifier si la requête a réussi
            if reponse.status_code != 200:
                print(f" Erreur de recherche : {reponse.status_code}")
                return []
            
            # Analyser le HTML
            soup = BeautifulSoup(reponse.text, 'html.parser')
            
            # Trouver les résultats de recherche
            resultats = soup.find_all('div', class_='g')
            
            # Créer une liste de résultats
            resultats_recherche = []
            
            # Limiter à 5 résultats
            for i, resultat in enumerate(resultats[:5], 1):
                # Extraire le titre
                titre_elem = resultat.find('h3')
                titre = titre_elem.get_text() if titre_elem else "Pas de titre"
                
                # Extraire le lien
                lien_elem = resultat.find('a')
                lien = lien_elem['href'] if lien_elem else "#"
                
                # Extraire l'extrait
                extrait_elem = resultat.find('div', class_='VwiC3b')
                extrait = extrait_elem.get_text() if extrait_elem else "Pas de description"
                
                # Ajouter le résultat à la liste
                resultats_recherche.append({
                    'title': titre[:100],
                    'link': lien,
                    'snippet': extrait[:200]
                })
            
            return resultats_recherche
        
        except Exception as e:
            print(f" Erreur lors de la recherche : {str(e)}")
            return []
    
    def _ajouter_mode_ecriture(self):
        """Ajoute un mode d'écriture assisté"""
        # Créer une fenêtre modale pour le mode écriture
        self.fenetre_ecriture = ctk.CTkToplevel(self)
        self.fenetre_ecriture.title("Mode Écriture Assisté")
        self.fenetre_ecriture.geometry("800x600")
        
        # Frame principale
        frame_principal = ctk.CTkFrame(self.fenetre_ecriture)
        frame_principal.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Options de type de document
        ctk.CTkLabel(frame_principal, text="Type de Document", font=("Arial", 12, "bold")).pack(pady=(10,5))
        
        self.var_type_document = tk.StringVar(value="Lettre")
        types_documents = ["Lettre", "Article", "Rapport", "Poème", "Histoire", "Email"]
        
        menu_type_document = ctk.CTkOptionMenu(
            frame_principal, 
            values=types_documents,
            variable=self.var_type_document,
            width=300
        )
        menu_type_document.pack(pady=5)
        
        # Zone de contexte
        ctk.CTkLabel(frame_principal, text="Contexte/Instructions", font=("Arial", 12, "bold")).pack(pady=(10,5))
        
        self.zone_contexte = ctk.CTkTextbox(
            frame_principal, 
            width=700, 
            height=100,
            font=("Arial", 12)
        )
        self.zone_contexte.pack(pady=5)
        
        # Zone de texte principale
        ctk.CTkLabel(frame_principal, text="Votre Texte", font=("Arial", 12, "bold")).pack(pady=(10,5))
        
        self.zone_texte = ctk.CTkTextbox(
            frame_principal, 
            width=700, 
            height=300,
            font=("Arial", 12)
        )
        self.zone_texte.pack(pady=5)
        
        # Boutons d'action
        frame_boutons = ctk.CTkFrame(frame_principal, fg_color="transparent")
        frame_boutons.pack(pady=10)
        
        # Bouton Générer
        bouton_generer = ctk.CTkButton(
            frame_boutons, 
            text="Générer avec IA", 
            command=self._generer_texte_ia,
            fg_color=COULEURS["accent_primaire"]
        )
        bouton_generer.pack(side='left', padx=5)
        
        # Bouton Enregistrer
        bouton_enregistrer = ctk.CTkButton(
            frame_boutons, 
            text="Enregistrer", 
            command=self._enregistrer_texte,
            fg_color=COULEURS["accent_secondaire"]
        )
        bouton_enregistrer.pack(side='left', padx=5)
        
        # Bouton Copier
        bouton_copier = ctk.CTkButton(
            frame_boutons, 
            text="Copier", 
            command=self._copier_texte,
            fg_color="gray"
        )
        bouton_copier.pack(side='left', padx=5)
    
    def _generer_texte_ia(self):
        """Génère du texte assisté par IA"""
        try:
            # Récupérer le contexte et le type de document
            contexte = self.zone_contexte.get("1.0", tk.END).strip()
            type_document = self.var_type_document.get()
            
            # Prompt pour Ollama
            prompt = f"""
            Tu es un assistant d'écriture professionnel. 
            Type de document : {type_document}
            Contexte : {contexte}
            
            Génère un texte adapté au type de document et au contexte fourni.
            """
            
            # Utiliser Ollama pour générer le texte
            reponse = self.gestionnaire_ollama.generer_reponse(prompt)
            
            # Afficher le texte généré
            self.zone_texte.delete("1.0", tk.END)
            self.zone_texte.insert(tk.END, reponse)
            
        except Exception as e:
            print(f" Erreur de génération de texte : {str(e)}")
    
    def _enregistrer_texte(self):
        """Enregistre le texte dans un fichier"""
        try:
            # Choisir le fichier de destination
            fichier = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[
                    ("Fichiers texte", "*.txt"),
                    ("Fichiers Word", "*.docx"),
                    ("Tous les fichiers", "*.*")
                ]
            )
            
            if fichier:
                texte = self.zone_texte.get("1.0", tk.END).strip()
                
                # Enregistrer selon l'extension
                if fichier.endswith('.docx'):
                    import docx
                    doc = docx.Document()
                    doc.add_paragraph(texte)
                    doc.save(fichier)
                else:
                    with open(fichier, 'w', encoding='utf-8') as f:
                        f.write(texte)
                
                print(f" Texte enregistré dans {fichier}")
        
        except Exception as e:
            print(f" Erreur d'enregistrement : {str(e)}")
    
    def _copier_texte(self):
        """Copie le texte dans le presse-papiers"""
        texte = self.zone_texte.get("1.0", tk.END).strip()
        self.clipboard_clear()
        self.clipboard_append(texte)
        print(" Le texte a été copié dans le presse-papiers")

    def _obtenir_modeles_ollama(self):
        """
        Récupère la liste des modèles Ollama disponibles.
        
        Returns:
            list: Liste des modèles Ollama installés
        """
        try:
            # Utiliser subprocess pour exécuter la commande Ollama
            import subprocess
            resultat = subprocess.run(
                ['ollama', 'list'], 
                capture_output=True, 
                text=True, 
                check=True
            )
            
            # Extraire les noms des modèles
            modeles = [ligne.split()[0] for ligne in resultat.stdout.strip().split('\n')[1:]]
            print(f" Modèles Ollama disponibles : {modeles}")
            return modeles
        
        except subprocess.CalledProcessError as e:
            print(f" Erreur lors de la récupération des modèles : {e}")
            return []
        except Exception as e:
            print(f" Erreur inattendue : {e}")
            return []

    def _telecharger_modele_ollama(self, modele):
        """
        Télécharge un modèle Ollama s'il n'est pas déjà installé.
        
        Args:
            modele (str): Nom du modèle à télécharger
        
        Returns:
            bool: True si le modèle est disponible, False sinon
        """
        try:
            import subprocess
            
            # Vérifier si le modèle existe déjà
            modeles_disponibles = self._obtenir_modeles_ollama()
            if modele in modeles_disponibles:
                print(f" Modèle {modele} déjà installé")
                return True
            
            # Télécharger le modèle
            print(f" Téléchargement du modèle {modele}...")
            resultat = subprocess.run(
                ['ollama', 'pull', modele], 
                capture_output=True, 
                text=True, 
                check=True
            )
            
            print(f" Modèle {modele} téléchargé avec succès")
            return True
        
        except subprocess.CalledProcessError as e:
            print(f" Erreur lors du téléchargement de {modele} : {e}")
            return False
        except Exception as e:
            print(f" Erreur inattendue lors du téléchargement : {e}")
            return False

    def _generer_reponse_ollama(self, prompt_systeme, modele='mistral', temperature=0.3):
        """
        Génère une réponse en utilisant le modèle Ollama spécifié.
        
        Args:
            prompt_systeme (str): Le prompt système détaillé
            modele (str, optional): Le modèle Ollama à utiliser. Défaut à 'mistral'.
            temperature (float, optional): La température pour la génération. Défaut à 0.3.
        
        Returns:
            str: La réponse générée par le modèle
        """
        try:
            # Utiliser le gestionnaire Ollama existant
            reponse = self.gestionnaire_ollama.chat(
                model=modele,
                messages=[
                    {
                        'role': 'system', 
                        'content': prompt_systeme
                    },
                    {
                        'role': 'user',
                        'content': 'Génère une réponse basée sur le prompt système.'
                    }
                ]
            )
            
            # Extraire le contenu de la réponse
            if isinstance(reponse, dict):
                resume = reponse.get('message', {}).get('content', '')
            elif isinstance(reponse, str):
                resume = reponse
            else:
                resume = "Impossible de générer une réponse."
            
            return resume.strip()
        
        except Exception as e:
            print(f" Erreur de génération Ollama : {e}")
            import traceback
            traceback.print_exc()
            return f"Erreur de génération : {str(e)}"

class InterfaceAssistantIA(ctk.CTk):
    def __init__(self, gestionnaire_config, gestionnaire_ollama):
        super().__init__()
        
        # Configuration de base
        self.title(" Assistant IA - Style ChatGPT")
        self.geometry("1600x900")
        
        # Centrer la fenêtre sur l'écran
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        x = (screen_width - 1600) // 2
        y = (screen_height - 900) // 2
        self.geometry(f'1600x900+{x}+{y}')
        
        # Configuration du thème
        ctk.set_appearance_mode("dark")  # Mode sombre par défaut
        ctk.set_default_color_theme("blue")  # Thème de couleur
        
        # Gestionnaire de configuration
        self.gestionnaire_config = gestionnaire_config
        
        # Gestionnaire Ollama
        self.gestionnaire_ollama = gestionnaire_ollama
        
        # Création de la structure principale
        self._creer_interface_principale()
        
    def _creer_interface_principale(self):
        """Crée l'interface principale avec système d'onglets"""
        # Configuration du grid
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        # Barre latérale pour l'historique
        self.barre_laterale = ctk.CTkFrame(
            self, 
            width=250, 
            fg_color=COULEURS["fond_secondaire"]
        )
        self.barre_laterale.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        
        # Bouton pour nouvelle conversation
        self.bouton_nouvelle_conv = ctk.CTkButton(
            self.barre_laterale, 
            text=" Nouvelle conversation", 
            command=self._creer_nouvelle_conversation,
            fg_color=COULEURS["accent_primaire"],
            hover_color=COULEURS["accent_secondaire"]
        )
        self.bouton_nouvelle_conv.pack(pady=10, padx=10, fill="x")
        
        # Liste des conversations historiques
        self.liste_historique = ctk.CTkScrollableFrame(
            self.barre_laterale, 
            fg_color="transparent"
        )
        self.liste_historique.pack(pady=10, padx=10, fill="both", expand=True)
        
        # Zone d'onglets
        self.gestionnaire_onglets = ctk.CTkTabview(
            self, 
            width=1350,  
            height=800,
            fg_color=COULEURS["fond_principal"]
        )
        self.gestionnaire_onglets.grid(row=0, column=1, sticky="nsew", padx=10, pady=10)
        
        # Créer un premier onglet par défaut
        self._creer_nouvelle_conversation()
        
    def _creer_nouvelle_conversation(self):
        """Crée une nouvelle conversation"""
        titre = f"Conversation {len(list(self.gestionnaire_onglets._tab_dict.keys())) + 1}"
        nouvel_onglet = self.gestionnaire_onglets.add(titre)
        
        # Créer l'onglet de conversation
        conversation_onglet = ConversationOnglet(
            nouvel_onglet, 
            self.gestionnaire_config,
            self.gestionnaire_ollama
        )
        conversation_onglet.pack(fill='both', expand=True)
        
        # Ajouter un bouton dans l'historique
        bouton_conv = ctk.CTkButton(
            self.liste_historique, 
            text=titre, 
            command=lambda t=titre: self.gestionnaire_onglets.set(t),
            fg_color=COULEURS["fond_secondaire"],
            hover_color=COULEURS["accent_primaire"]
        )
        bouton_conv.pack(pady=5, fill="x")
        
    def executer(self):
        """Lance l'application"""
        self.mainloop()